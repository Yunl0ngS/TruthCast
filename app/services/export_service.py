from __future__ import annotations

from datetime import datetime, timezone
from html import escape
from io import BytesIO
from typing import Any

from app.schemas.export import ExportDataRequest

_RISK_LABEL_MAP = {
    "credible": "可信",
    "suspicious": "可疑",
    "high_risk": "高风险",
    "needs_context": "需要补充语境",
    "likely_misinformation": "疑似不实信息",
}

_RISK_LEVEL_MAP = {
    "low": "低",
    "medium": "中",
    "high": "高",
    "critical": "严重",
}

_STANCE_MAP = {
    "support": "支持",
    "supportive": "支持",
    "oppose": "反对",
    "opposing": "反对",
    "refute": "反驳",
    "insufficient": "证据不足",
    "insufficient_evidence": "证据不足",
    "doubt": "质疑",
    "mixed": "混合",
    "neutral": "中立",
}

_SCENARIO_MAP = {
    "general": "通用",
    "health": "医疗健康",
    "governance": "政务治理",
    "security": "公共安全",
    "media": "媒体传播",
    "technology": "科技产业",
    "education": "教育校园",
}

_DOMAIN_MAP = {
    "health": "医疗健康",
    "governance": "政务治理",
    "security": "公共安全",
    "media": "媒体传播",
    "technology": "科技产业",
    "education": "教育校园",
    "general": "通用",
}

_SOURCE_TYPE_MAP = {
    "local_kb": "本地知识库",
    "web_live": "联网检索",
    "web_summary": "联网聚合",
}

_EMOTION_MAP = {
    "anger": "愤怒",
    "fear": "恐惧",
    "sadness": "悲伤",
    "surprise": "惊讶",
    "neutral": "中性",
    "joy": "喜悦",
    "disgust": "厌恶",
    "anticipation": "期待",
    "trust": "信任",
}

_SIM_STANCE_MAP = {
    "supportive": "支持",
    "opposing": "反对",
    "neutral": "中立",
    "skeptical": "质疑",
    "mixed": "混合",
    "dismissive": "否定",
    "curious": "好奇",
}

_PRIORITY_MAP = {"urgent": "紧急", "high": "高", "medium": "中"}
_CATEGORY_MAP = {
    "official": "官方",
    "media": "媒体",
    "platform": "平台",
    "user": "用户",
}


def _zh(value: str | None, mapping: dict[str, str]) -> str:
    if not value:
        return "-"
    return mapping.get(value, value)


def _safe(value: str | None) -> str:
    return value or "-"


def _exported_at(data: ExportDataRequest) -> str:
    if data.exported_at:
        return data.exported_at
    return datetime.now(timezone.utc).isoformat()


def _percent(value: float | int) -> str:
    return f"{float(value) * 100:.1f}%"


def _zh_stance(value: str | None) -> str:
    mapped = _zh(value, _SIM_STANCE_MAP)
    if mapped == (value or "-"):
        mapped = _zh(value, _STANCE_MAP)
    return mapped


def _collect_primary_clarification(content_data: object) -> object | None:
    if content_data is None:
        return None
    clarification = getattr(content_data, "clarification", None)
    if clarification is not None:
        return clarification
    clarifications = getattr(content_data, "clarifications", None)
    primary_id = getattr(content_data, "primary_clarification_id", None)
    if not clarifications:
        return None
    if primary_id:
        for item in clarifications:
            item_id = (
                item.get("id") if isinstance(item, dict) else getattr(item, "id", None)
            )
            if item_id == primary_id:
                content = (
                    item.get("content")
                    if isinstance(item, dict)
                    else getattr(item, "content", None)
                )
                if content is not None:
                    return content
    latest = sorted(
        clarifications,
        key=lambda x: (
            x.get("generated_at", "")
            if isinstance(x, dict)
            else getattr(x, "generated_at", "")
        ),
        reverse=True,
    )[0]
    return (
        latest.get("content")
        if isinstance(latest, dict)
        else getattr(latest, "content", None)
    )


def _clarification_field(clarification: object, field: str) -> str:
    if isinstance(clarification, dict):
        return str(clarification.get(field, "") or "")
    return str(getattr(clarification, field, "") or "")


def _apply_word_zh_font(document: Any, font_name: str = "Microsoft YaHei") -> None:
    try:
        from docx.oxml.ns import qn
    except Exception:
        return

    style_names = ["Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]
    for style_name in style_names:
        try:
            style = document.styles[style_name]
        except Exception:
            continue
        style.font.name = font_name
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run._element.get_or_add_rPr().rFonts.set(
                            qn("w:eastAsia"), font_name
                        )


def _group_evidence(data: ExportDataRequest) -> list[tuple[str, str, list[object]]]:
    if data.report and data.report.claim_reports:
        grouped: list[tuple[str, str, list[object]]] = []
        for claim_report in data.report.claim_reports:
            grouped.append(
                (
                    claim_report.claim.claim_id,
                    claim_report.claim.claim_text,
                    list(claim_report.evidences),
                )
            )
        return grouped

    claim_text_map = {claim.claim_id: claim.claim_text for claim in data.claims}
    buckets: dict[str, list[object]] = {}
    for evidence in data.evidences:
        buckets.setdefault(evidence.claim_id or "unknown", []).append(evidence)

    grouped = []
    for claim_id, items in buckets.items():
        grouped.append(
            (claim_id, claim_text_map.get(claim_id, "未关联到具体主张"), items)
        )
    return grouped


def _build_html(data: ExportDataRequest) -> str:
    parts: list[str] = []
    parts.append("<h1>TruthCast 智能研判报告</h1>")
    parts.append(f"<p><strong>导出时间：</strong>{escape(_exported_at(data))}</p>")

    parts.append("<h2>原始输入</h2>")
    parts.append(f"<div class='quote'>{escape(data.input_text)}</div>")

    if data.report and (
        data.report.source_url
        or data.report.source_title
        or data.report.source_publish_date
    ):
        parts.append("<h2>原始来源</h2>")
        parts.append("<table><tr><th>项目</th><th>值</th></tr>")
        parts.append(
            f"<tr><td>标题</td><td>{escape(_safe(data.report.source_title))}</td></tr>"
        )
        source_url = _safe(data.report.source_url)
        if source_url != "-":
            parts.append(
                f"<tr><td>链接</td><td><a href='{escape(source_url)}'>{escape(source_url)}</a></td></tr>"
            )
        else:
            parts.append("<tr><td>链接</td><td>-</td></tr>")
        parts.append(
            f"<tr><td>发布时间</td><td>{escape(_safe(data.report.source_publish_date))}</td></tr>"
        )
        parts.append("</table>")

    if data.detect_data:
        detect = data.detect_data
        parts.append("<h2>风险初判</h2>")
        parts.append("<table><tr><th>项目</th><th>值</th></tr>")
        parts.append(
            f"<tr><td>风险标签</td><td>{escape(_zh(detect.label, _RISK_LABEL_MAP))}</td></tr>"
        )
        parts.append(f"<tr><td>风险分数</td><td>{detect.score}</td></tr>")
        parts.append(f"<tr><td>置信度</td><td>{detect.confidence}</td></tr>")
        parts.append("</table>")
        if detect.reasons:
            parts.append("<h3>风险理由</h3><ul>")
            for reason in detect.reasons:
                parts.append(f"<li>{escape(reason)}</li>")
            parts.append("</ul>")

    if data.claims:
        parts.append("<h2>主张抽取</h2>")
        parts.append("<table><tr><th>ID</th><th>主张内容</th></tr>")
        for claim in data.claims:
            parts.append(
                "<tr>"
                f"<td>{escape(claim.claim_id)}</td>"
                f"<td>{escape(claim.claim_text)}</td>"
                "</tr>"
            )
        parts.append("</table>")

    grouped_evidence = _group_evidence(data)
    if grouped_evidence:
        parts.append("<h2>证据链</h2>")
        for claim_id, claim_text, evidences in grouped_evidence:
            parts.append(f"<h3>{escape(claim_id)}: {escape(claim_text)}</h3>")
            if not evidences:
                parts.append("<p>暂无对齐证据</p>")
                continue
            for idx, evidence in enumerate(evidences, start=1):
                is_summary = getattr(evidence, "source_type", "") == "web_summary"
                title = (
                    _safe(getattr(evidence, "summary", None))
                    if is_summary
                    else _safe(getattr(evidence, "title", None))
                )
                parts.append(f"<h4>证据 {idx}: {escape(title)}</h4>")
                parts.append("<table><tr><th>属性</th><th>值</th></tr>")
                parts.append(
                    f"<tr><td>立场</td><td>{escape(_zh(getattr(evidence, 'stance', None), _STANCE_MAP))}</td></tr>"
                )
                parts.append(
                    f"<tr><td>来源</td><td>{escape(_safe(getattr(evidence, 'source', None)))}</td></tr>"
                )
                parts.append(
                    f"<tr><td>来源类型</td><td>{escape(_zh(getattr(evidence, 'source_type', None), _SOURCE_TYPE_MAP))}</td></tr>"
                )
                parts.append(
                    f"<tr><td>权重</td><td>{float(getattr(evidence, 'source_weight', 0.0)):.2f}</td></tr>"
                )
                parts.append(
                    f"<tr><td>领域</td><td>{escape(_zh(getattr(evidence, 'domain', None), _DOMAIN_MAP))}</td></tr>"
                )
                if getattr(evidence, "alignment_confidence", None) is not None:
                    parts.append(
                        f"<tr><td>对齐置信度</td><td>{float(getattr(evidence, 'alignment_confidence')):.2f}</td></tr>"
                    )
                parts.append("</table>")

                summary = getattr(evidence, "summary", None)
                if summary and not is_summary:
                    parts.append(f"<p><strong>摘要：</strong>{escape(summary)}</p>")

                rationale = getattr(evidence, "alignment_rationale", None)
                if rationale:
                    parts.append(
                        f"<p><strong>对齐理由：</strong>{escape(rationale)}</p>"
                    )

                source_urls = getattr(evidence, "source_urls", None)
                if source_urls:
                    parts.append(
                        f"<p><strong>来源链接（{len(source_urls)}条）</strong></p><ul>"
                    )
                    for url in source_urls:
                        parts.append(
                            f"<li><a href='{escape(url)}'>{escape(url)}</a></li>"
                        )
                    parts.append("</ul>")
                else:
                    url = _safe(getattr(evidence, "url", None))
                    parts.append(
                        f"<p><strong>链接：</strong><a href='{escape(url)}'>{escape(url)}</a></p>"
                    )

    if data.report:
        report = data.report
        parts.append("<h2>综合报告</h2>")
        domains_zh = (
            "、".join(_zh(d, _DOMAIN_MAP) for d in report.evidence_domains)
            if report.evidence_domains
            else "-"
        )
        parts.append("<table><tr><th>项目</th><th>值</th></tr>")
        parts.append(
            f"<tr><td>风险评级</td><td>{escape(_zh(report.risk_label, _RISK_LABEL_MAP))}（{escape(_zh(report.risk_level, _RISK_LEVEL_MAP))}风险）</td></tr>"
        )
        parts.append(f"<tr><td>风险分数</td><td>{report.risk_score}</td></tr>")
        parts.append(
            f"<tr><td>识别场景</td><td>{escape(_zh(report.detected_scenario, _SCENARIO_MAP))}</td></tr>"
        )
        parts.append(f"<tr><td>证据覆盖域</td><td>{escape(domains_zh)}</td></tr>")
        parts.append("</table>")
        parts.append(f"<p><strong>摘要：</strong>{escape(report.summary)}</p>")
        if report.suspicious_points:
            parts.append("<h3>可疑点</h3><ul>")
            for point in report.suspicious_points:
                parts.append(f"<li>{escape(point)}</li>")
            parts.append("</ul>")
        if report.claim_reports:
            parts.append("<h3>主张级结论</h3>")
            for claim_report in report.claim_reports:
                parts.append(f"<h4>{escape(claim_report.claim.claim_id)}</h4>")
                parts.append(
                    f"<p><strong>主张：</strong>{escape(claim_report.claim.claim_text)}</p>"
                )
                parts.append(
                    f"<p><strong>最终立场：</strong>{escape(_zh_stance(claim_report.final_stance))}</p>"
                )
                if claim_report.notes:
                    parts.append("<ul>")
                    for note in claim_report.notes:
                        parts.append(f"<li>{escape(note)}</li>")
                    parts.append("</ul>")

    if data.simulation:
        simulation = data.simulation
        parts.append("<h2>舆情预演</h2>")

        parts.append("<h3>情绪分布</h3>")
        parts.append("<table><tr><th>情绪</th><th>占比</th></tr>")
        for key, value in simulation.emotion_distribution.items():
            parts.append(
                f"<tr><td>{escape(_zh(key, _EMOTION_MAP))}</td><td>{_percent(value)}</td></tr>"
            )
        parts.append("</table>")

        parts.append("<h3>立场分布</h3>")
        parts.append("<table><tr><th>立场</th><th>占比</th></tr>")
        for key, value in simulation.stance_distribution.items():
            stance_text = _zh(key, _SIM_STANCE_MAP)
            if stance_text == key:
                stance_text = _zh(key, _STANCE_MAP)
            parts.append(
                f"<tr><td>{escape(stance_text)}</td><td>{_percent(value)}</td></tr>"
            )
        parts.append("</table>")

        if simulation.narratives:
            parts.append("<h3>叙事分支</h3>")
            for idx, narrative in enumerate(simulation.narratives, start=1):
                parts.append(f"<h4>{idx}. {escape(narrative.title)}</h4>")
                parts.append(
                    f"<p><strong>概率：</strong>{_percent(narrative.probability)}</p>"
                )
                parts.append(
                    f"<p><strong>立场：</strong>{escape(_zh_stance(narrative.stance))}</p>"
                )
                parts.append(
                    f"<p><strong>触发词：</strong>{escape(', '.join(narrative.trigger_keywords) if narrative.trigger_keywords else '-')}</p>"
                )
                parts.append(
                    f"<p><strong>代表言论：</strong>{escape(narrative.sample_message)}</p>"
                )

        if simulation.timeline:
            parts.append("<h3>时间线</h3>")
            parts.append("<table><tr><th>小时</th><th>事件</th><th>预估触达</th></tr>")
            for item in simulation.timeline:
                parts.append(
                    f"<tr><td>{item.hour}</td><td>{escape(item.event)}</td><td>{escape(item.expected_reach)}</td></tr>"
                )
            parts.append("</table>")

        if simulation.flashpoints:
            parts.append("<h3>引爆点</h3><ul>")
            for flashpoint in simulation.flashpoints:
                parts.append(f"<li>{escape(flashpoint)}</li>")
            parts.append("</ul>")

        parts.append("<h3>应对建议</h3>")
        if simulation.suggestion.summary:
            parts.append(
                f"<p><strong>{escape(simulation.suggestion.summary)}</strong></p>"
            )
        if simulation.suggestion.actions:
            parts.append(
                "<table><tr><th>优先级</th><th>类别</th><th>行动</th><th>时间</th><th>责任方</th></tr>"
            )
            for action in simulation.suggestion.actions:
                parts.append(
                    "<tr>"
                    f"<td>{escape(_zh(action.priority, _PRIORITY_MAP))}</td>"
                    f"<td>{escape(_zh(action.category, _CATEGORY_MAP))}</td>"
                    f"<td>{escape(action.action)}</td>"
                    f"<td>{escape(_safe(action.timeline))}</td>"
                    f"<td>{escape(_safe(action.responsible))}</td>"
                    "</tr>"
                )
            parts.append("</table>")

    if data.content:
        content = data.content
        parts.append("<h2>公关响应</h2>")
        primary = _collect_primary_clarification(content)
        if primary is not None:
            parts.append("<h3>澄清稿（主稿）</h3>")
            parts.append(
                f"<p><strong>短版：</strong>{escape(_clarification_field(primary, 'short'))}</p>"
            )
            parts.append(
                f"<p><strong>中版：</strong>{escape(_clarification_field(primary, 'medium'))}</p>"
            )
            parts.append(
                f"<p><strong>长版：</strong>{escape(_clarification_field(primary, 'long'))}</p>"
            )

        if content.faq:
            parts.append("<h3>FAQ</h3><ol>")
            for item in content.faq:
                parts.append(
                    f"<li><strong>{escape(item.question)}</strong><br />{escape(item.answer)}</li>"
                )
            parts.append("</ol>")

        if content.platform_scripts:
            parts.append("<h3>多平台话术</h3>")
            for script in content.platform_scripts:
                platform = (
                    script.platform.value
                    if hasattr(script.platform, "value")
                    else str(script.platform)
                )
                parts.append(f"<h4>{escape(platform)}</h4>")
                parts.append(f"<p>{escape(script.content)}</p>")

    body = "\n".join(parts)
    return f"""
<!doctype html>
<html lang=\"zh-CN\">
<head>
  <meta charset=\"utf-8\" />
  <style>
    @page {{ size: A4; margin: 18mm; }}
    body {{ font-family: 'Noto Sans CJK SC', 'Microsoft YaHei', 'PingFang SC', sans-serif; font-size: 12px; line-height: 1.6; color: #1f2937; }}
    h1 {{ font-size: 24px; margin: 0 0 12px; }}
    h2 {{ font-size: 18px; margin: 20px 0 8px; border-bottom: 1px solid #e5e7eb; padding-bottom: 4px; }}
    h3 {{ font-size: 14px; margin: 14px 0 6px; }}
    h4 {{ font-size: 12px; margin: 10px 0 4px; }}
    p, li {{ white-space: pre-wrap; word-break: break-word; }}
    table {{ width: 100%; border-collapse: collapse; margin: 8px 0 12px; }}
    th, td {{ border: 1px solid #d1d5db; text-align: left; vertical-align: top; padding: 6px; }}
    .quote {{ background: #f9fafb; border-left: 4px solid #93c5fd; padding: 8px 10px; }}
  </style>
</head>
<body>
{body}
<hr />
<p>本报告由 TruthCast 智能研判台自动生成，仅供辅助决策参考。</p>
</body>
</html>
"""


def generate_pdf_bytes(data: ExportDataRequest) -> bytes:
    try:
        from weasyprint import HTML
    except Exception:
        return _generate_pdf_with_reportlab(data)

    html = _build_html(data)
    try:
        rendered = HTML(string=html).write_pdf()
        if not rendered:
            return _generate_pdf_with_reportlab(data)
        return bytes(rendered)
    except Exception:
        return _generate_pdf_with_reportlab(data)


def _generate_pdf_with_reportlab(data: ExportDataRequest) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas
    except Exception as exc:
        raise RuntimeError(
            "PDF 导出失败：weasyprint 不可用且 reportlab 未安装，请安装依赖后重试"
        ) from exc

    html_text = _build_html(data)
    lines = [
        line.strip()
        for line in html_text.replace("<", "\n<").replace(">", ">\n").splitlines()
    ]
    lines = [line for line in lines if line and not line.startswith("<")]

    width, height = A4
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    font_name = "Helvetica"
    line_height = 16

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        font_name = "STSong-Light"
        line_height = 18
    except Exception:
        font_name = "Helvetica"

    y = height - 30
    pdf.setFont(font_name, 11)
    for line in lines:
        text = line.replace("&nbsp;", " ")
        if len(text) > 110:
            chunks = [text[i : i + 110] for i in range(0, len(text), 110)]
        else:
            chunks = [text]
        for chunk in chunks:
            y -= line_height
            if y < 30:
                pdf.showPage()
                pdf.setFont(font_name, 11)
                y = height - 30
            pdf.drawString(30, y, chunk)

    pdf.save()
    return buffer.getvalue()


def generate_word_bytes(data: ExportDataRequest) -> bytes:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Word 导出依赖未安装（python-docx）") from exc

    doc = Document()
    doc.add_heading("TruthCast 智能研判报告", level=1)
    doc.add_paragraph(f"导出时间：{_exported_at(data)}")

    doc.add_heading("原始输入", level=2)
    doc.add_paragraph(data.input_text)

    if data.report and (
        data.report.source_url
        or data.report.source_title
        or data.report.source_publish_date
    ):
        doc.add_heading("原始来源", level=2)
        src_table = doc.add_table(rows=3, cols=2)
        src_table.rows[0].cells[0].text = "标题"
        src_table.rows[0].cells[1].text = _safe(data.report.source_title)
        src_table.rows[1].cells[0].text = "链接"
        src_table.rows[1].cells[1].text = _safe(data.report.source_url)
        src_table.rows[2].cells[0].text = "发布时间"
        src_table.rows[2].cells[1].text = _safe(data.report.source_publish_date)

    if data.detect_data:
        detect = data.detect_data
        doc.add_heading("风险初判", level=2)
        table = doc.add_table(rows=4, cols=2)
        table.rows[0].cells[0].text = "风险标签"
        table.rows[0].cells[1].text = _zh(detect.label, _RISK_LABEL_MAP)
        table.rows[1].cells[0].text = "风险分数"
        table.rows[1].cells[1].text = str(detect.score)
        table.rows[2].cells[0].text = "置信度"
        table.rows[2].cells[1].text = str(detect.confidence)
        table.rows[3].cells[0].text = "风险理由"
        table.rows[3].cells[1].text = (
            "；".join(detect.reasons) if detect.reasons else "-"
        )

    if data.claims:
        doc.add_heading("主张抽取", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "ID"
        table.rows[0].cells[1].text = "主张内容"
        for claim in data.claims:
            row = table.add_row().cells
            row[0].text = claim.claim_id
            row[1].text = claim.claim_text

    grouped_evidence = _group_evidence(data)
    if grouped_evidence:
        doc.add_heading("证据链", level=2)
        for claim_id, claim_text, evidences in grouped_evidence:
            doc.add_heading(f"{claim_id}: {claim_text}", level=3)
            if not evidences:
                doc.add_paragraph("暂无对齐证据")
                continue
            for idx, evidence in enumerate(evidences, start=1):
                is_summary = getattr(evidence, "source_type", "") == "web_summary"
                title = (
                    _safe(getattr(evidence, "summary", None))
                    if is_summary
                    else _safe(getattr(evidence, "title", None))
                )
                doc.add_heading(f"证据 {idx}: {title}", level=4)
                detail = doc.add_table(rows=6, cols=2)
                detail.rows[0].cells[0].text = "立场"
                detail.rows[0].cells[1].text = _zh(
                    getattr(evidence, "stance", None), _STANCE_MAP
                )
                detail.rows[1].cells[0].text = "来源"
                detail.rows[1].cells[1].text = _safe(getattr(evidence, "source", None))
                detail.rows[2].cells[0].text = "来源类型"
                detail.rows[2].cells[1].text = _zh(
                    getattr(evidence, "source_type", None), _SOURCE_TYPE_MAP
                )
                detail.rows[3].cells[0].text = "权重"
                detail.rows[3].cells[
                    1
                ].text = f"{float(getattr(evidence, 'source_weight', 0.0)):.2f}"
                detail.rows[4].cells[0].text = "领域"
                detail.rows[4].cells[1].text = _zh(
                    getattr(evidence, "domain", None), _DOMAIN_MAP
                )
                detail.rows[5].cells[0].text = "对齐置信度"
                align_conf = getattr(evidence, "alignment_confidence", None)
                detail.rows[5].cells[1].text = (
                    f"{float(align_conf):.2f}" if align_conf is not None else "-"
                )

                rationale = getattr(evidence, "alignment_rationale", None)
                if rationale:
                    doc.add_paragraph(f"对齐理由：{rationale}")
                source_urls = getattr(evidence, "source_urls", None)
                if source_urls:
                    doc.add_paragraph("来源链接：")
                    for url in source_urls:
                        doc.add_paragraph(url, style="List Bullet")
                else:
                    doc.add_paragraph(f"链接：{_safe(getattr(evidence, 'url', None))}")

    if data.report:
        report = data.report
        doc.add_heading("综合报告", level=2)
        domains_zh = (
            "、".join(_zh(d, _DOMAIN_MAP) for d in report.evidence_domains)
            if report.evidence_domains
            else "-"
        )
        doc.add_paragraph(
            f"风险评级：{_zh(report.risk_label, _RISK_LABEL_MAP)}（{_zh(report.risk_level, _RISK_LEVEL_MAP)}风险）"
        )
        doc.add_paragraph(f"风险分数：{report.risk_score}")
        doc.add_paragraph(f"识别场景：{_zh(report.detected_scenario, _SCENARIO_MAP)}")
        doc.add_paragraph(f"证据覆盖域：{domains_zh}")
        doc.add_paragraph(f"摘要：{report.summary}")
        if report.suspicious_points:
            doc.add_paragraph("可疑点：")
            for point in report.suspicious_points:
                doc.add_paragraph(point, style="List Bullet")
        if report.claim_reports:
            doc.add_heading("主张级结论", level=3)
            for claim_report in report.claim_reports:
                doc.add_paragraph(
                    f"{claim_report.claim.claim_id}: {claim_report.claim.claim_text}"
                )
                doc.add_paragraph(f"最终立场：{_zh_stance(claim_report.final_stance)}")
                for note in claim_report.notes:
                    doc.add_paragraph(note, style="List Bullet")

    if data.simulation:
        simulation = data.simulation
        doc.add_heading("舆情预演", level=2)

        doc.add_heading("情绪分布", level=3)
        emotion_table = doc.add_table(rows=1, cols=2)
        emotion_table.rows[0].cells[0].text = "情绪"
        emotion_table.rows[0].cells[1].text = "占比"
        for key, value in simulation.emotion_distribution.items():
            row = emotion_table.add_row().cells
            row[0].text = _zh(key, _EMOTION_MAP)
            row[1].text = _percent(value)

        doc.add_heading("立场分布", level=3)
        stance_table = doc.add_table(rows=1, cols=2)
        stance_table.rows[0].cells[0].text = "立场"
        stance_table.rows[0].cells[1].text = "占比"
        for key, value in simulation.stance_distribution.items():
            row = stance_table.add_row().cells
            mapped = _zh(key, _SIM_STANCE_MAP)
            if mapped == key:
                mapped = _zh(key, _STANCE_MAP)
            row[0].text = mapped
            row[1].text = _percent(value)

        if simulation.narratives:
            doc.add_heading("叙事分支", level=3)
            for idx, item in enumerate(simulation.narratives, start=1):
                doc.add_heading(f"{idx}. {item.title}", level=4)
                doc.add_paragraph(f"概率：{_percent(item.probability)}")
                doc.add_paragraph(f"立场：{_zh_stance(item.stance)}")
                doc.add_paragraph(
                    f"触发词：{', '.join(item.trigger_keywords) if item.trigger_keywords else '-'}"
                )
                doc.add_paragraph(f"代表言论：{item.sample_message}")

        if simulation.timeline:
            doc.add_heading("时间线", level=3)
            timeline_table = doc.add_table(rows=1, cols=3)
            timeline_table.rows[0].cells[0].text = "小时"
            timeline_table.rows[0].cells[1].text = "事件"
            timeline_table.rows[0].cells[2].text = "预估触达"
            for item in simulation.timeline:
                row = timeline_table.add_row().cells
                row[0].text = str(item.hour)
                row[1].text = item.event
                row[2].text = item.expected_reach

        if simulation.flashpoints:
            doc.add_heading("引爆点", level=3)
            for flashpoint in simulation.flashpoints:
                doc.add_paragraph(flashpoint, style="List Bullet")

        doc.add_heading("应对建议", level=3)
        if simulation.suggestion.summary:
            doc.add_paragraph(simulation.suggestion.summary)
        if simulation.suggestion.actions:
            action_table = doc.add_table(rows=1, cols=5)
            action_table.rows[0].cells[0].text = "优先级"
            action_table.rows[0].cells[1].text = "类别"
            action_table.rows[0].cells[2].text = "行动"
            action_table.rows[0].cells[3].text = "时间"
            action_table.rows[0].cells[4].text = "责任方"
            for action in simulation.suggestion.actions:
                row = action_table.add_row().cells
                row[0].text = _zh(action.priority, _PRIORITY_MAP)
                row[1].text = _zh(action.category, _CATEGORY_MAP)
                row[2].text = action.action
                row[3].text = _safe(action.timeline)
                row[4].text = _safe(action.responsible)

    if data.content:
        content = data.content
        doc.add_heading("公关响应", level=2)
        primary = _collect_primary_clarification(content)
        if primary is not None:
            short_text = _clarification_field(primary, "short")
            medium_text = _clarification_field(primary, "medium")
            long_text = _clarification_field(primary, "long")
            doc.add_heading("澄清稿（主稿）", level=3)
            doc.add_paragraph(f"短版：{short_text}")
            doc.add_paragraph(f"中版：{medium_text}")
            doc.add_paragraph(f"长版：{long_text}")
        if content.faq:
            doc.add_heading("FAQ", level=3)
            for item in content.faq:
                doc.add_paragraph(f"Q: {item.question}")
                doc.add_paragraph(f"A: {item.answer}")
        if content.platform_scripts:
            doc.add_heading("多平台话术", level=3)
            for script in content.platform_scripts:
                platform = (
                    script.platform.value
                    if hasattr(script.platform, "value")
                    else str(script.platform)
                )
                doc.add_heading(platform, level=4)
                doc.add_paragraph(script.content)

    _apply_word_zh_font(doc)
    doc.add_paragraph("本报告由 TruthCast 智能研判台自动生成，仅供辅助决策参考。")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
